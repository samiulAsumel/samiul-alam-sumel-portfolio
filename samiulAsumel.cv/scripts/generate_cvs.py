#!/usr/bin/env python3
"""Generate Samiul_Alam_Sumel_CV.docx and Samiul_Alam_Sumel_Freelance_CV.docx.

Content is sourced strictly from the portfolio site (../index.html) -- same
facts, same honest framing, no sugarcoating, no claims beyond what real
GitHub projects and self-directed study actually support.

Usage: python3 scripts/generate_cvs.py
Requires: pip install python-docx
Output files are written to the repo root (one level up from this script).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).resolve().parent.parent

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GOLD = RGBColor(0xB8, 0x86, 0x2A)
TEXT = RGBColor(0x2A, 0x2A, 0x2A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
LINK = RGBColor(0x1E, 0x3A, 0x5F)

FONT = "Calibri"


def add_hyperlink(paragraph, url, text, color=LINK, underline=True, bold=False, size=None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)

    if bold:
        rPr.append(OxmlElement("w:b"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    rPr.append(c)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9.5)
    style.font.color.rgb = TEXT
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.28)
    sec.bottom_margin = Inches(0.2)
    sec.left_margin = Inches(0.52)
    sec.right_margin = Inches(0.52)

    lb = doc.styles["List Bullet"]
    lb.font.name = FONT
    lb.paragraph_format.space_before = Pt(0)
    lb.paragraph_format.space_after = Pt(1)
    lb.paragraph_format.line_spacing = 1.0


def para_border_bottom(paragraph, color="B8862A", sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_name_header(doc, name, title, contact_parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(name.upper())
    r.font.size = Pt(19)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    r2.font.size = Pt(11.5)
    r2.font.bold = True
    r2.font.color.rgb = GOLD

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(8)
    for i, (label, url) in enumerate(contact_parts):
        if i > 0:
            sep = p3.add_run("   |   ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = MUTED
        if url:
            add_hyperlink(p3, url, label, color=NAVY, underline=False, size=9)
        else:
            r = p3.add_run(label)
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED
    para_border_bottom(p3, color="1E3A5F", sz=10)


def add_section_heading(doc, text, space_before=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = NAVY
    para_border_bottom(p, color="B8862A", sz=6)
    return p


def add_body(doc, text, space_after=4, size=9.5, italic=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_lead=None, space_after=1, size=9.3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = TEXT
        text = text[len(bold_lead):] if text.startswith(bold_lead) else text
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.color.rgb = TEXT
    return p


def add_role_header(doc, role, org, period, subnote=None):
    # NOTE: left-aligned (not justified) -- uses a right tab stop to place
    # the date; justify would break that layout.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(role)
    r1.bold = True
    r1.font.size = Pt(10.3)
    r1.font.color.rgb = NAVY
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.4), alignment=3)  # right align tab
    r_tab = p.add_run("\t" + period)
    r_tab.font.size = Pt(9.2)
    r_tab.italic = True
    r_tab.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(org)
    r2.bold = True
    r2.font.size = Pt(9.7)
    r2.font.color.rgb = GOLD
    if subnote:
        r3 = p2.add_run("  —  " + subnote)
        r3.font.size = Pt(9.2)
        r3.italic = True
        r3.font.color.rgb = MUTED


def add_skill_line(doc, category, items, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(category + ":  ")
    r1.bold = True
    r1.font.size = Pt(9.6)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(items)
    r2.font.size = Pt(9.6)
    r2.font.color.rgb = TEXT


def add_project_entry(doc, name, links, problem, solution, tags):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(10.1)
    r.font.color.rgb = NAVY
    if links:
        sep = p.add_run("   —   ")
        sep.font.size = Pt(8.8)
        sep.font.color.rgb = MUTED
        for i, (label, url) in enumerate(links):
            if i > 0:
                s = p.add_run("  |  ")
                s.font.size = Pt(8.8)
                s.font.color.rgb = MUTED
            add_hyperlink(p, url, label, color=GOLD, underline=True, size=8.8)

    add_body(doc, "Problem: " + problem, space_after=1, size=9.1, italic=True, color=MUTED)
    add_body(doc, solution, space_after=1, size=9.3)
    add_body(doc, "Tech: " + tags, space_after=3, size=8.6, italic=True, color=MUTED)


def add_compact_project(doc, name, links, sentence, tag_note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(9.1)
    r.font.color.rgb = NAVY
    if tag_note:
        rn = p.add_run("  (" + tag_note + ")")
        rn.italic = True
        rn.font.size = Pt(8)
        rn.font.color.rgb = MUTED
    if links:
        sep = p.add_run("   —   ")
        sep.font.size = Pt(8.2)
        sep.font.color.rgb = MUTED
        for i, (label, url) in enumerate(links):
            if i > 0:
                s = p.add_run("  |  ")
                s.font.size = Pt(8.2)
                s.font.color.rgb = MUTED
            add_hyperlink(p, url, label, color=GOLD, underline=True, size=8.2)
    add_body(doc, sentence, space_after=1, size=8.5)


CONTACT = [
    ("Mongla, Bangladesh", None),
    ("sa.sumel91@gmail.com", "mailto:sa.sumel91@gmail.com"),
    ("+880 1312 312 512", None),
    ("sasumel.pages.dev", "https://sasumel.pages.dev/"),
    ("github.com/samiulAsumel", "https://github.com/samiulAsumel"),
    ("linkedin.com/in/samiul-alam-sumel", "https://linkedin.com/in/samiul-alam-sumel"),
]

SKILLS = [
    ("Domain & Port Operations", "Port billing & tariff systems, wharfrent/VAT/levy calculation, "
                                  "cargo & vessel tracking, C&F agent workflows, "
                                  "customs/regulatory compliance context"),
    ("Frontend & JavaScript", "HTML5 (Semantic), CSS3, Responsive Design, "
                               "Tailwind CSS, DaisyUI, JavaScript (ES6+), DOM & Events, "
                               "Fetch & Promises / Async-Await, Array/Object Methods, ES Modules"),
    ("React & Next.js", "Own: React (Hooks, Context), React Router, React Hook Form. "
                         "AI-assisted: TanStack Query, Axios (incl. Interceptors), "
                         "Next.js App Router, NextAuth"),
    ("Backend, Data & Deployment", "Own: Git & GitHub (Branching, PRs), deployment "
                                    "(Cloudflare Pages/Firebase/Netlify). AI-assisted: "
                                    "Node.js & Express, REST APIs (CRUD), MongoDB (Atlas), "
                                    "Aggregation Pipeline, JWT Auth, Firebase Auth & Admin "
                                    "SDK, Next.js API Routes, Stripe Payments, "
                                    "Pagination & Search"),
    ("Linux & Automation (self-directed)", "Bash Scripting, Linux Administration (Ubuntu), "
                                            "systemd, Cron & Timers, Monitoring & Alerting Scripts"),
]

PROJECTS = [
    dict(
        name='Port Billing Calculator',
        links=[("Live", "https://portbill.pages.dev"), ("GitHub", "https://github.com/samiulAsumel/portbill")],
        problem="Port billing system creates permanent entries — C&F agents had no way to "
                "estimate wharfrent charges before committing, causing billing disputes, "
                "excessive dwell time, and repeated counter visits.",
        solution="Real-time advance wharfrent calculator: slab-based charge computation, VAT "
                 "and levy calculation, inside/outside cargo split, hoisting charge "
                 "auto-calculation, print-ready A4 output. Actively used by C&F agents at "
                 "Mongla Port daily.",
        tags="JavaScript, HTML/CSS, Cloudflare Pages, PWA, Print Output",
    ),
    dict(
        name='Daily Car Balance & Location Tracking System ("carview")',
        links=[("Live", "https://carview.pages.dev"), ("GitHub", "https://github.com/samiulAsumel/carview")],
        problem="Vehicle positions across warehouse, shed, and yard were recorded in one "
                "person's personal Excel file — invisible to all other staff, creating "
                "bottlenecks and dependency on a single individual.",
        solution="Offline-first PWA tracking 8 port locations, with Cloudflare Worker + private "
                 "GitHub repo sync, a Chart.js analytics dashboard (7 charts + KPIs), 13 report "
                 "sections, and Excel export. Built for Mongla Port Authority's Traffic Department.",
        tags="PWA, Cloudflare Worker, Chart.js, GitHub Sync, Excel Export",
    ),
    dict(
        name="OT Bill Management System",
        links=[("Live", "https://otbill.pages.dev"), ("GitHub", "https://github.com/samiulAsumel/otbill")],
        problem="Overtime billing was a fully manual Excel process — multi-step hourly rate "
                "calculations and final bill generation done by hand every cycle.",
        solution="Staff enters an employee profile once; the system generates the complete "
                 "final OT bill instantly with correct hourly rate, cumulative date-wise OT "
                 "calculation, and A4 print-ready output. Reduced billing cycle time from "
                 "hours to minutes.",
        tags="JavaScript, Employee DB, Auto-Calc, A4 Print, Cloudflare Pages",
    ),
    dict(
        name="Client Intake Form",
        links=[("Live", "https://clif91.pages.dev/"), ("GitHub", "https://github.com/samiulAsumel/client-intake-form")],
        problem="Collecting project requirements from a client over email or chat is "
                "unstructured — details get missed, wasting time on both sides.",
        solution="Single-file, client-side project requirement intake form — structured "
                 "questions, no backend required, mailto-based report so the completed brief "
                 "lands directly in an inbox.",
        tags="JavaScript, HTML/CSS, Client-Side Only",
    ),
]

ADDITIONAL_PROJECTS = [
    dict(
        name="World Kitchen Atlas",
        tag_note=None,
        links=[("Live", "https://kitchenatlas.pages.dev"),
               ("GitHub", "https://github.com/samiulAsumel/world-kitchen-atlas")],
        sentence="Next.js static-export culinary encyclopedia on Cloudflare Pages, with a "
                 "Cloudflare Worker proxying a private GitHub data repo, admin CRUD, and a "
                 "Durable-Object-backed visit counter.",
    ),
    dict(
        name="Outreach Copilot",
        tag_note=None,
        links=[("Live", "https://outreach-copilot.pages.dev"),
               ("GitHub", "https://github.com/samiulAsumel/outreach-copilot")],
        sentence="Single-user AI-drafted cold-outreach tool — Cloudflare Pages + Workers API, "
                 "D1-backed leads/analytics, Workers AI draft generation across email/LinkedIn/"
                 "WhatsApp/cover letter; every send stays a manual, deliberate step.",
    ),
    dict(
        name="SalahSync",
        tag_note=None,
        links=[("Live", "https://salahsync.pages.dev"),
               ("GitHub", "https://github.com/samiulAsumel/salahsync")],
        sentence="Offline-first installable PWA that auto-generates a full daily prayer/routine "
                 "schedule from Fajr time and office hours — offline solar prayer-time "
                 "calculation, streak tracking, bilingual Bangla/English UI.",
    ),
    dict(
        name="JARVIS",
        tag_note="source available, no public demo",
        links=[("GitHub", "https://github.com/samiulAsumel/jarvis")],
        sentence="Self-hosted personal AI assistant with multi-model routing (Groq, OpenRouter, "
                 "Claude, Gemini, local Ollama), voice I/O, and persistent memory.",
    ),
]

INFRA_NOTE = (
    "Self-directed practice building production-style Bash automation for deployment "
    "(auto-deploy-pipeline), SSL certificate renewal (auto-ssl-manager), and log/security "
    "monitoring (log-analyzer-alert). Daily Linux (Ubuntu) user; self-driven learning, not "
    "job-critical production work."
)

STUDY_SUMMARY = (
    "Self-paced study for code comprehension, not a certificate — starting from The Odin "
    "Project's Intermediate HTML and CSS course, then Full Stack Open (University of "
    "Helsinki): JavaScript fundamentals, Node.js/Express in depth, testing, and Git "
    "workflow, built from assignment briefs rather than a course guide."
)


def build_final_cv():
    doc = Document()
    set_base_style(doc)

    add_name_header(doc, "MD Samiul Alam Sumel", "AI Product Engineer — Port & Logistics Domain", CONTACT)

    add_section_heading(doc, "Professional Summary", space_before=0)
    add_body(
        doc,
        "AI Product Engineer with 12+ years of hands-on experience in cargo billing, "
        "vessel/vehicle tracking, and port authority workflows at Mongla Port Authority. "
        "Finds the operational problem, designs the system, builds it with AI-assisted "
        "development, then rebuilds each piece independently from memory before counting "
        "it as learned. 3 self-built tools — billing, vehicle tracking, payroll automation "
        "— are in daily production use. Skills backed by that shipped software and "
        "self-directed study (The Odin Project, Full Stack Open), not a bootcamp "
        "certificate. Builds frontend independently; backend/data work with AI-assisted "
        "development (Claude Code), closed through a self-authored engineering roadmap. "
        "Not a bootcamp graduate or senior engineer, but bringing 12+ years of professional "
        "discipline from a demanding operational job. Seeking a remote AI Product Engineer, "
        "Port Operations Technologist, Logistics Systems Engineer, or full-stack role, "
        "worldwide, with no time zone constraints.",
        size=9.4,
    )

    add_section_heading(doc, "Technical Skills")
    for cat, items in SKILLS:
        add_skill_line(doc, cat, items)

    add_section_heading(doc, "Professional Experience")
    add_role_header(
        doc,
        "Senior Outdoor Assistant — Port Operations & IT Systems",
        "Mongla Port Authority",
        "Nov 2017 – Present",
        "Bangladesh's second-largest international seaport",
    )
    add_body(doc, "Core Port Operations", space_after=1, size=9.3, italic=True, color=GOLD, align=None)
    for t in [
        "Manage end-to-end wharfrent billing cycle — slab-wise charge computation, VAT and "
        "levy calculation, inside/outside cargo split, and final billing issuance for active "
        "C&F agents; identified the recurring manual-calculation bottleneck and independently "
        "built and deployed portbill, now used daily by the billing desk.",
        "Track cargo and vehicle positions across warehouse, shed, and yard terminal "
        "locations in real time; the single-person, Excel-only version of this record was "
        "the direct trigger for building carview, now the Traffic Department's shared, "
        "version-controlled tracking system.",
        "Operate the port automation billing system end-to-end and coordinate C&F agent "
        "documentation and advance payment planning.",
    ]:
        add_bullet(doc, t)
    add_body(doc, "Self-Initiated Web Development & Learning", space_after=1, size=9.3, italic=True, color=GOLD, align=None)
    for t in [
        "Independently identified critical inefficiencies in manual paper-based billing, "
        "overtime payroll calculation, and vehicle tracking workflows — designed and "
        "deployed 3 live web applications now used daily by C&F agents and port staff, "
        "built with plain JavaScript/HTML/CSS and AI-assisted development as a productivity "
        "tool.",
        "Worked through a large part of a full-stack curriculum (HTML/CSS, JavaScript, "
        "React, part of Node.js/Express and Next.js) alongside a full-time port operations "
        "career, before stepping away to focus on shipping real projects.",
        "Maintain an active public GitHub portfolio (github.com/samiulAsumel) demonstrating "
        "continuous technical growth.",
    ]:
        add_bullet(doc, t)

    add_role_header(
        doc,
        "Junior Outdoor Assistant — Port & Terminal Operations",
        "Mongla Port Authority",
        "Nov 2013 – Nov 2017",
    )
    for t in [
        "Assisted in port inspections, vessel operations, and cargo handling under senior "
        "supervision; monitored cargo dwell time and vehicle positioning across terminal "
        "locations.",
        "Built foundational expertise in wharfrent billing, maritime logistics, and customs "
        "clearance that underpins all later self-initiated digital work.",
    ]:
        add_bullet(doc, t)

    add_section_heading(doc, "Featured Projects")
    for pr in PROJECTS:
        add_project_entry(doc, pr["name"], pr["links"], pr["problem"], pr["solution"], pr["tags"])
    for pr in ADDITIONAL_PROJECTS:
        add_compact_project(doc, pr["name"], pr["links"], pr["sentence"], pr["tag_note"])

    add_section_heading(doc, "Linux / Automation Practice", space_before=6)
    add_body(doc, INFRA_NOTE, space_after=3, size=8.8)

    add_section_heading(doc, "Self-Directed Study")
    add_role_header(
        doc,
        "The Odin Project & Full Stack Open",
        "Self-Directed — In Progress",
        "Aug 2026 – Present",
    )
    add_body(doc, STUDY_SUMMARY, space_after=4, size=9.5)

    add_section_heading(doc, "Education & Languages")
    add_skill_line(doc, "Higher Secondary Certificate — Business Studies",
                   "Khulna Public College, 2009 – 2011", space_after=2)
    add_skill_line(doc, "Secondary School Certificate — Business Studies",
                   "Mongla Bandar Secondary School, 1999 – 2009", space_after=2)
    add_skill_line(doc, "Languages", "Bengali — Native   |   English — Professional", space_after=0)

    path = OUT_DIR / "Samiul_Alam_Sumel_CV.docx"
    doc.save(path)
    print("Saved:", path)


def build_freelance_cv():
    doc = Document()
    set_base_style(doc)

    add_name_header(doc, "MD Samiul Alam Sumel", "AI Product Engineer — Port & Logistics Domain — Freelance & Remote", CONTACT)

    add_section_heading(doc, "How I Work", space_before=0)
    add_body(
        doc,
        "Send the workflow and the problem — I identify what's actually going wrong, using "
        "12+ years of port operations experience, then design, build, and deploy the "
        "solution: frontend independently, backend/data work with AI-assisted development "
        "(Claude Code). Every build gets a second pass — rebuilt from memory, no AI, then "
        "explained back before I call it done. An AI Product Engineer (React/Next.js, "
        "Node.js/Express, MongoDB) who worked through a large part of an intensive full-stack "
        "curriculum, backed by 3 live web apps built for real daily use — not demos. "
        "Available for freelance and "
        "contract work worldwide, no time zone constraints.",
        size=9.4,
    )

    add_section_heading(doc, "Services I Deliver")
    services = [
        ("Full-Stack Web Apps",
         "React/Next.js on the frontend, Node.js/Express + MongoDB on the backend — REST "
         "APIs, JWT/Firebase authentication, and Stripe payments when the job needs them.",
         "3 apps in daily production use (vanilla JS + Cloudflare Workers); React/Next.js "
         "via World Kitchen Atlas, Node/MongoDB backend via AI-assisted development."),
        ("Billing & Tracking Dashboards",
         "Custom billing calculators, tracking dashboards, and operational tools — the same "
         "kind of tool built to solve real problems at my own workplace.",
         "Built for real daily use, not a demo."),
        ("Port & Logistics Workflow Digitisation",
         "Turning paper-based or Excel-based port and logistics workflows — wharfrent "
         "billing, overtime payroll, cargo/vehicle tracking — into deployed, team-wide web "
         "tools, backed by first-hand operational knowledge.",
         "Proven with 3 tools converted from manual/Excel processes."),
    ]
    for title, desc, proof in services:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.3)
        r.font.color.rgb = NAVY
        add_body(doc, desc, space_after=1, size=9.6)
        add_body(doc, "✓ " + proof, space_after=3, size=9.2, color=GOLD, italic=True, align=None)

    add_section_heading(doc, "Remote & Freelance Readiness")
    for t in [
        "No time zone constraints — any schedule, worldwide.",
        "Self-managed: scopes, builds, and deploys independently, manually verified before "
        "every release.",
        "Async-friendly English, Git/GitHub-based collaboration (branching, PRs).",
        "Every past project is publicly verifiable before you hire.",
    ]:
        add_bullet(doc, t)

    add_section_heading(doc, "Featured Projects")
    for pr in PROJECTS:
        add_project_entry(doc, pr["name"], pr["links"], pr["problem"], pr["solution"], pr["tags"])

    add_section_heading(doc, "Technical Skills")
    for cat, items in SKILLS:
        add_skill_line(doc, cat, items)

    add_section_heading(doc, "Professional Background")
    add_body(
        doc,
        "12+ years of professional experience at Mongla Port Authority — Senior Outdoor "
        "Assistant, Port Operations & IT Systems (Nov 2017 – Present); Junior Outdoor "
        "Assistant, Port & Terminal Operations (Nov 2013 – Nov 2017). Wharfrent billing, "
        "terminal operations, cargo dwell time management, and C&F agent coordination — a "
        "track record of discipline and reliability, not just technical skill.",
        size=9.6, space_after=6,
    )

    add_section_heading(doc, "Self-Directed Study")
    add_role_header(
        doc,
        "The Odin Project & Full Stack Open",
        "Self-Directed — In Progress",
        "Aug 2026 – Present",
    )
    add_body(doc, STUDY_SUMMARY, space_after=6, size=9.5)

    add_section_heading(doc, "Education & Languages")
    add_skill_line(doc, "Education",
                   "HSC — Business Studies, Khulna Public College (2009–2011); "
                   "SSC — Business Studies, Mongla Bandar Secondary School (1999–2009)",
                   space_after=2)
    add_skill_line(doc, "Languages", "Bengali — Native   |   English — Professional", space_after=2)

    path = OUT_DIR / "Samiul_Alam_Sumel_Freelance_CV.docx"
    doc.save(path)
    print("Saved:", path)


if __name__ == "__main__":
    build_final_cv()
    build_freelance_cv()
